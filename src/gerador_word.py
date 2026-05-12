import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def adicionar_titulo(documento, texto):
    titulo = documento.add_heading(texto, level=1)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER


def adicionar_subtitulo(documento, texto):
    documento.add_heading(texto, level=2)


def adicionar_paragrafo(documento, texto):
    paragrafo = documento.add_paragraph()
    paragrafo.add_run(texto)
    paragrafo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def adicionar_tabela(documento, df, titulo):
    adicionar_subtitulo(documento, titulo)

    if df is None or df.empty:
        adicionar_paragrafo(documento, "Não há dados disponíveis para esta seção.")
        return

    tabela = documento.add_table(rows=1, cols=len(df.columns))
    tabela.style = "Table Grid"

    hdr_cells = tabela.rows[0].cells

    for i, coluna in enumerate(df.columns):
        hdr_cells[i].text = str(coluna)

    for _, linha in df.iterrows():
        row_cells = tabela.add_row().cells
        for i, valor in enumerate(linha):
            row_cells[i].text = str(valor)

    documento.add_paragraph("Fonte: ENAMED Performance Analytics.")


def adicionar_figura(documento, caminho_figura, legenda):
    if caminho_figura and os.path.exists(caminho_figura):
        documento.add_picture(caminho_figura, width=Inches(6))
        paragrafo = documento.add_paragraph(legenda)
        paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER


def gerar_texto_resumo(estatisticas, config):
    instituicao = config.get("instituicao", "Instituição de Ensino Superior")

    return (
        f"O presente relatório apresenta a análise dos resultados de simulados e/ou provas "
        f"do ENAMED aplicados aos estudantes de Medicina da {instituicao}. "
        f"A base analisada contemplou {estatisticas['total_estudantes']} estudantes. "
        f"A média geral de acertos foi de {estatisticas['media_acertos']} questões, "
        f"com mediana de {estatisticas['mediana_acertos']}, desvio-padrão de "
        f"{estatisticas['desvio_padrao_acertos']} e média percentual de "
        f"{estatisticas['media_percentual']}%. Esses indicadores permitem uma leitura "
        f"institucional do desempenho acadêmico, subsidiando ações do NDE, da coordenação "
        f"do curso, do colegiado e dos processos de melhoria contínua."
    )


def gerar_interpretacao_turmas(tabela_turma):
    if tabela_turma.empty:
        return "Não foram identificados dados suficientes para análise por turma."

    melhor = tabela_turma.sort_values("media_percentual", ascending=False).iloc[0]
    menor = tabela_turma.sort_values("media_percentual", ascending=True).iloc[0]

    return (
        f"A análise por turma evidencia variação no desempenho médio entre os grupos avaliados. "
        f"A turma com maior média percentual foi {melhor['turma']}, com "
        f"{melhor['media_percentual']}% de acertos. A menor média percentual foi observada "
        f"na turma {menor['turma']}, com {menor['media_percentual']}% de acertos. "
        f"Esses resultados permitem identificar desigualdades formativas, necessidades de "
        f"intervenção pedagógica e oportunidades de acompanhamento longitudinal."
    )


def gerar_interpretacao_ciclos(tabela_ciclo):
    if tabela_ciclo.empty:
        return "Não foram identificados dados suficientes para análise por ciclo formativo."

    melhor = tabela_ciclo.sort_values("media_percentual", ascending=False).iloc[0]

    return (
        f"A análise por ciclos formativos permite observar a progressão cognitiva dos estudantes "
        f"ao longo da formação médica. O ciclo com maior média percentual foi "
        f"{melhor['ciclo']}, com {melhor['media_percentual']}% de acertos. "
        f"Essa leitura contribui para avaliar a coerência entre matriz curricular, desenvolvimento "
        f"de competências e desempenho nos eixos avaliativos do ENAMED."
    )


def gerar_interpretacao_areas(tabela_areas):
    if tabela_areas is None or tabela_areas.empty:
        return (
            "A análise por grandes áreas não foi realizada porque as colunas específicas "
            "das áreas do ENAMED não foram identificadas na planilha."
        )

    maior = tabela_areas.sort_values("media_percentual", ascending=False).iloc[0]
    menor = tabela_areas.sort_values("media_percentual", ascending=True).iloc[0]

    return (
        f"A análise por grandes áreas do ENAMED demonstrou maior desempenho médio em "
        f"{maior['area']}, com {maior['media_percentual']}% de acertos, enquanto a menor "
        f"média foi observada em {menor['area']}, com {menor['media_percentual']}%. "
        f"Esses achados permitem direcionar estratégias de reforço, revisão curricular, "
        f"monitoramento de competências e planejamento de intervenções pedagógicas específicas."
    )


def gerar_interpretacao_risco(tabela_risco):
    if tabela_risco.empty:
        return "Não foram identificados estudantes em risco."

    contagem = tabela_risco["classificacao_risco"].value_counts().to_dict()

    moderado = contagem.get("Risco pedagógico moderado", 0)
    critico = contagem.get("Risco pedagógico crítico", 0)

    return (
        f"A análise de risco pedagógico identificou {moderado} estudantes em risco moderado "
        f"e {critico} estudantes em risco crítico. Essa classificação considera o desempenho "
        f"individual em relação à média da própria turma, permitindo uma leitura contextualizada "
        f"do rendimento acadêmico e favorecendo ações de tutoria, monitoria, revisão de conteúdos "
        f"e acompanhamento pedagógico individualizado."
    )


def gerar_plano_acao():
    return (
        "Recomenda-se que os resultados sejam apresentados e discutidos em reunião do NDE "
        "e do colegiado do curso, com registro em ata. As turmas com menor desempenho relativo "
        "devem receber planejamento específico de reforço pedagógico, enquanto os estudantes "
        "classificados em risco devem ser acompanhados por estratégias de tutoria, monitoria "
        "e orientação acadêmica. As áreas com menor desempenho médio devem subsidiar ações "
        "de revisão curricular, oficinas de aprendizagem, análise de itens e alinhamento entre "
        "competências previstas no PPC e práticas avaliativas."
    )


def gerar_word(resultados, caminhos_graficos, pasta_execucao, config):
    pasta_word = os.path.join(pasta_execucao, "word")
    os.makedirs(pasta_word, exist_ok=True)

    caminho_word = os.path.join(
        pasta_word,
        "relatorio_enamed_performance_analytics.docx"
    )

    documento = Document()

    adicionar_titulo(documento, config.get("nome_relatorio", "Relatório ENAMED Performance Analytics"))

    adicionar_paragrafo(
        documento,
        "Sistema institucional para análise de desempenho em simulados e provas do ENAMED em escolas médicas."
    )

    adicionar_subtitulo(documento, "1. Resumo executivo")
    adicionar_paragrafo(
        documento,
        gerar_texto_resumo(resultados["estatisticas_gerais"], config)
    )

    adicionar_subtitulo(documento, "2. Estatísticas gerais")
    adicionar_tabela(
        documento,
        __import__("pandas").DataFrame([resultados["estatisticas_gerais"]]),
        "Tabela 1 – Estatísticas gerais do desempenho"
    )

    adicionar_subtitulo(documento, "3. Desempenho por turma")
    adicionar_paragrafo(
        documento,
        gerar_interpretacao_turmas(resultados["analise_por_turma"])
    )
    adicionar_tabela(
        documento,
        resultados["analise_por_turma"],
        "Tabela 2 – Indicadores de desempenho por turma"
    )

    if caminhos_graficos.get("media_turma"):
        adicionar_figura(
            documento,
            caminhos_graficos["media_turma"][0],
            "Figura 1 – Média percentual de acertos por turma."
        )

    adicionar_subtitulo(documento, "4. Desempenho por ciclo formativo")
    adicionar_paragrafo(
        documento,
        gerar_interpretacao_ciclos(resultados["analise_por_ciclo"])
    )
    adicionar_tabela(
        documento,
        resultados["analise_por_ciclo"],
        "Tabela 3 – Indicadores de desempenho por ciclo formativo"
    )

    if caminhos_graficos.get("media_ciclo"):
        adicionar_figura(
            documento,
            caminhos_graficos["media_ciclo"][0],
            "Figura 2 – Média percentual de acertos por ciclo formativo."
        )

    adicionar_subtitulo(documento, "5. Grandes áreas do ENAMED")
    adicionar_paragrafo(
        documento,
        gerar_interpretacao_areas(resultados["analise_por_areas"])
    )
    adicionar_tabela(
        documento,
        resultados["analise_por_areas"],
        "Tabela 4 – Desempenho por grandes áreas do ENAMED"
    )

    if caminhos_graficos.get("areas"):
        adicionar_figura(
            documento,
            caminhos_graficos["areas"][0],
            "Figura 3 – Desempenho médio por grande área do ENAMED."
        )

    adicionar_subtitulo(documento, "6. Estudantes em risco pedagógico")
    adicionar_paragrafo(
        documento,
        gerar_interpretacao_risco(resultados["estudantes_em_risco"])
    )
    adicionar_tabela(
        documento,
        resultados["estudantes_em_risco"],
        "Tabela 5 – Estudantes classificados por risco pedagógico"
    )

    if caminhos_graficos.get("risco"):
        adicionar_figura(
            documento,
            caminhos_graficos["risco"][0],
            "Figura 4 – Distribuição dos estudantes por classificação de risco."
        )

    adicionar_subtitulo(documento, "7. Bonificação acadêmica")
    adicionar_paragrafo(
        documento,
        "A bonificação acadêmica foi calculada automaticamente com base no desempenho percentual acima da média da própria turma, conforme as regras parametrizadas no arquivo config.json."
    )
    adicionar_tabela(
        documento,
        resultados["bonificacao"],
        "Tabela 6 – Bonificação acadêmica por estudante"
    )

    adicionar_subtitulo(documento, "8. Plano de ação pedagógico")
    adicionar_paragrafo(documento, gerar_plano_acao())

    adicionar_subtitulo(documento, "9. Encaminhamentos institucionais")
    adicionar_paragrafo(
        documento,
        "Os resultados produzidos pelo ENAMED Performance Analytics devem ser utilizados como evidência institucional para planejamento pedagógico, acompanhamento da aprendizagem, análise de progressão formativa, qualificação das avaliações internas e apoio aos processos de gestão acadêmica baseada em evidências."
    )

    documento.save(caminho_word)

    return caminho_word